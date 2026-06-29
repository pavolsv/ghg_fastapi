import json
import re
from pathlib import Path
from typing import Any

from fastapi import APIRouter, Request, Depends, HTTPException
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse, FileResponse
from fastapi.templating import Jinja2Templates
from sqlmodel import Session, col, select
from datetime import datetime

from audit_log import add_change_log
from database import engine
from model import Device, EmissionFactor604, EmissionRecord, DataChangeLog, GWPReference
from constants.refrigerant_factors import get_rate_by_code
from routers.emission_source import _derive_device_code_display

router = APIRouter(prefix="/calculation", tags=["calculation"])
templates = Jinja2Templates(directory="templates")


TEMPLATE_TAG_PATTERN = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")

REPORT_EDITOR_OUTLINE: list[dict[str, Any]] = [
    {"key": "chapter_1", "title": "第1章 公司簡介與政策聲明", "level": 1, "editable": False},
    {"key": "chapter_1_1_preface", "title": "1.1 前言", "level": 2, "editable": True},
    {"key": "chapter_1_2_expected_use", "title": "1.2 預期用途", "level": 2, "editable": True},
    {"key": "chapter_1_3_company_profile", "title": "1.3 公司簡介", "level": 2, "editable": False},
    {"key": "chapter_1_3_1_basic_info", "title": "1.3.1 公司基本資料", "level": 3, "editable": True},
    {"key": "chapter_1_3_2_history", "title": "1.3.2 企業發展沿革", "level": 3, "editable": True},
    {"key": "chapter_1_3_3_main_business", "title": "1.3.3 主要營業項目", "level": 3, "editable": True},
    {"key": "chapter_1_3_4_operations", "title": "1.3.4 營運概況", "level": 3, "editable": True},
    {"key": "chapter_2", "title": "第2章 盤查邊界設定", "level": 1, "editable": False},
    {"key": "chapter_2_1_org_boundary", "title": "2.1 組織邊界設定", "level": 2, "editable": True},
    {"key": "chapter_2_2_operational_boundary", "title": "2.2 營運邊界", "level": 2, "editable": True},
    {"key": "chapter_3", "title": "第3章 報告溫室氣體排放量", "level": 1, "editable": False},
    {"key": "chapter_3_1_emission_types", "title": "3.1 溫室氣體排放類型與排放量說明", "level": 2, "editable": True},
    {"key": "chapter_3_2_category1", "title": "3.2 直接溫室氣體排放（類別1）", "level": 2, "editable": True},
    {"key": "chapter_3_3_category2", "title": "3.3 能源間接溫室氣體排放（類別2）", "level": 2, "editable": True},
    {"key": "chapter_3_4_total", "title": "3.4 溫室氣體總排放量", "level": 2, "editable": True},
    {"key": "chapter_4", "title": "第4章 數據品質管理", "level": 1, "editable": False},
    {"key": "chapter_4_1_quantification_method", "title": "4.1 量化方法", "level": 2, "editable": False},
    {"key": "chapter_4_1_1_activity_collection", "title": "4.1.1 活動數據蒐集與轉換方式", "level": 3, "editable": True},
    {"key": "chapter_4_1_2_conversion_and_gwp", "title": "4.1.2 排放係數與 GWP 換算依據", "level": 3, "editable": True},
    {"key": "chapter_4_2_method_change", "title": "4.2 量化方法變更說明", "level": 2, "editable": True},
    {"key": "chapter_4_3_factor_change", "title": "4.3 排放係數與變更說明", "level": 2, "editable": True},
    {"key": "chapter_4_4_significant_digits", "title": "4.4 有效位數", "level": 2, "editable": False},
    {"key": "chapter_4_4_1_activity_data", "title": "4.4.1 活動數據", "level": 3, "editable": True},
    {"key": "chapter_4_4_2_emission_factors", "title": "4.4.2 排放係數", "level": 3, "editable": True},
    {"key": "chapter_4_4_3_lhv", "title": "4.4.3 低位熱值", "level": 3, "editable": True},
    {"key": "chapter_4_4_4_unit_emission", "title": "4.4.4 每單位體積或重量之溫室氣體排放量", "level": 3, "editable": True},
    {"key": "chapter_4_4_5_single_gas", "title": "4.4.5 單一排放源之各溫室氣體排放量", "level": 3, "editable": True},
    {"key": "chapter_4_4_6_single_co2e", "title": "4.4.6 單一排放源之單一溫室氣體排放當量", "level": 3, "editable": True},
    {"key": "chapter_4_4_7_source_total", "title": "4.4.7 單一排放源之總排放當量", "level": 3, "editable": True},
    {"key": "chapter_4_4_8_grand_total", "title": "4.4.8 總排放當量彙總", "level": 3, "editable": True},
    {"key": "chapter_4_5_major_source_flow", "title": "4.5 重大排放源之資訊流", "level": 2, "editable": True},
    {"key": "chapter_5", "title": "第5章 基準年", "level": 1, "editable": False},
    {"key": "chapter_5_1_base_year", "title": "5.1 基準年設定", "level": 2, "editable": True},
    {"key": "chapter_6", "title": "第6章 參考文獻", "level": 1, "editable": False},
    {"key": "chapter_6_references", "title": "6.1 參考文獻", "level": 2, "editable": True},
    {"key": "appendix", "title": "附件", "level": 1, "editable": False},
    {"key": "appendix_1_org_chart", "title": "附件一 公司組織圖", "level": 2, "editable": True},
    {"key": "appendix_2_location", "title": "附件二 公司地理位置", "level": 2, "editable": True},
    {"key": "appendix_3_layout", "title": "附件三 公司廠區配置圖", "level": 2, "editable": True},
    {"key": "appendix_4_info_flow", "title": "附件四 重大排放源之資訊流", "level": 2, "editable": True},
    {"key": "appendix_5_tools", "title": "附件五 溫室氣體盤查工具", "level": 2, "editable": True},
]

REPORT_EDITOR_EDITABLE_KEYS = [
    item["key"] for item in REPORT_EDITOR_OUTLINE if item.get("editable")
]


def _build_default_report_sections(snapshot: dict[str, Any]) -> dict[str, str]:
    company = snapshot.get("company_summary", {}) if isinstance(snapshot.get("company_summary"), dict) else {}
    summary = snapshot.get("summary", {}) if isinstance(snapshot.get("summary"), dict) else {}
    scope_counts = snapshot.get("scope_source_counts", [])
    emission_type_totals = snapshot.get("emission_type_totals", [])
    top_devices = snapshot.get("top_devices", [])

    company_name = str(company.get("company_name") or "本公司")
    company_address = str(company.get("address") or "（請補公司地址）")
    owner = str(company.get("owner") or "（請補負責人）")
    tax_id = str(company.get("tax_id") or "（請補統編）")
    inventory_year = str(snapshot.get("inventory_year") or datetime.now().year)

    total_co2e = float(summary.get("total_co2e") or 0.0)
    record_count = int(summary.get("record_count") or 0)

    scope_lines = []
    if isinstance(scope_counts, list):
        for item in scope_counts:
            if not isinstance(item, dict):
                continue
            scope_name = str(item.get("scope") or "未分類")
            count = int(item.get("source_count") or 0)
            scope_lines.append(f"- {scope_name}：{count} 項")
    if not scope_lines:
        scope_lines.append("- 尚無可用範疇統計，請確認排放源與設備資料。")

    emission_lines = []
    if isinstance(emission_type_totals, list):
        for item in emission_type_totals[:5]:
            if not isinstance(item, dict):
                continue
            emission_lines.append(
                f"- {item.get('emission_type', '未分類')}：{item.get('total_co2e', 0)} kg CO₂e"
            )
    if not emission_lines:
        emission_lines.append("- 尚無可用排放類型統計資料。")

    top_device_lines: list[str] = []
    if isinstance(top_devices, list):
        for item in top_devices[:5]:
            if not isinstance(item, dict):
                continue
            top_device_lines.append(
                f"- {item.get('device_name', '未命名設備')}：{item.get('total_co2e', 0)} kg CO₂e"
            )
    if not top_device_lines:
        top_device_lines.append("- 尚無設備排放量排序資料。")

    return {
        "chapter_1_1_preface": (
            "本報告書依據 ISO 14064-1 與主管機關盤查指引編製，"
            "用以揭露組織於盤查年度之溫室氣體排放現況與管理作法。"
        ),
        "chapter_1_2_expected_use": (
            "本報告書可作為內部減碳管理、供應鏈揭露、外部查證及利害關係人溝通之依據。"
        ),
        "chapter_1_3_company_profile": "",
        "chapter_1_3_1_basic_info": (
            f"事業名稱：{company_name}\n"
            f"事業地址：{company_address}\n"
            f"事業負責人：{owner}\n"
            f"統一編號：{tax_id}"
        ),
        "chapter_1_3_2_history": (
            f"{company_name} 可於此補充公司成立背景、發展沿革與關鍵里程碑。"
        ),
        "chapter_1_3_3_main_business": "請補充主要產品、服務項目與供應鏈角色。",
        "chapter_1_3_4_operations": "請補充廠區/辦公據點、員工規模與營運概況。",
        "chapter_2_1_org_boundary": (
            "本次盤查以控制權法界定組織邊界，"
            f"盤查年度為 {inventory_year} 年，範圍為公司實際營運據點。"
        ),
        "chapter_2_2_operational_boundary": "\n".join(scope_lines),
        "chapter_3_1_emission_types": "\n".join(emission_lines),
        "chapter_3_2_category1": "請說明類別1（直接排放）之主要排放源、活動數據、排放係數與計算依據。",
        "chapter_3_3_category2": "請說明類別2（能源間接排放）之外購能源使用量、排放係數與計算結果。",
        "chapter_3_4_total": (
            f"本公司 {inventory_year} 年溫室氣體總排放量為 {round(total_co2e, 3)} kg CO₂e，"
            f"本期共納入 {record_count} 筆排放紀錄。"
        ),
        "chapter_4_1_quantification_method": "",
        "chapter_4_1_1_activity_collection": (
            "各排放源活動數據由權責部門依既有單據、系統或設備銘牌蒐集，"
            "並依盤查單位需求進行單位換算與年度彙整。"
        ),
        "chapter_4_1_2_conversion_and_gwp": (
            "排放量計算以排放係數法為主，並依所採用之 GWP 版本將各氣體轉換為 CO₂e。"
        ),
        "chapter_4_2_method_change": "本年度量化方法如有異動，請補充變更原因、差異分析與回溯調整說明。",
        "chapter_4_3_factor_change": "本年度排放係數如有更新，請補充係數來源、版本與影響評估。",
        "chapter_4_4_significant_digits": "",
        "chapter_4_4_1_activity_data": "活動數據原則可揭露至小數點後第 4 位，第 5 位四捨五入。",
        "chapter_4_4_2_emission_factors": "排放係數採主管機關公告或經內部核可之來源，依原公告位數使用。",
        "chapter_4_4_3_lhv": "低位熱值若有使用，建議揭露至小數點後第 2 位。",
        "chapter_4_4_4_unit_emission": "每單位體積或重量之溫室氣體排放量，應依計算公式與來源條件統一位數規則。",
        "chapter_4_4_5_single_gas": "單一排放源之各溫室氣體排放量可揭露至小數點後第 4 位。",
        "chapter_4_4_6_single_co2e": "單一排放源之單一溫室氣體排放當量可揭露至小數點後第 4 位。",
        "chapter_4_4_7_source_total": "單一排放源總排放當量可揭露至小數點後第 4 位。",
        "chapter_4_4_8_grand_total": "總排放當量彙總建議揭露至小數點後第 3 位，第 4 位四捨五入。",
        "chapter_4_5_major_source_flow": "\n".join(top_device_lines),
        "chapter_5_1_base_year": (
            f"本公司基準年設定為 {inventory_year} 年，"
            f"基準年總排放量為 {round(total_co2e, 3)} kg CO₂e。"
        ),
        "chapter_6_references": (
            "1. 溫室氣體排放係數管理表 6.0.4 版\n"
            "2. 溫室氣體排放量盤查作業指引 2022.5\n"
            "3. 經濟部能源局電力排碳係數公告\n"
            "4. ISO 14064-1:2018"
        ),
        "appendix_1_org_chart": "請插入公司組織圖。",
        "appendix_2_location": "請插入公司地理位置圖與地址說明。",
        "appendix_3_layout": "請插入廠區配置圖。",
        "appendix_4_info_flow": "請補充重大排放源資訊流圖與說明。",
        "appendix_5_tools": "請補充盤查工具或系統輸出佐證。",
    }


def _ensure_report_sections(payload: dict[str, Any], snapshot: dict[str, Any]) -> dict[str, Any]:
    normalized = dict(payload or {})
    report_sections = normalized.get("report_sections")
    if not isinstance(report_sections, dict):
        report_sections = {}

    compatibility_map = {
        "chapter_1_3_1_basic_info": "chapter_1_3_company_profile",
        "chapter_4_1_1_activity_collection": "chapter_4_1_quantification_method",
        "chapter_4_4_1_activity_data": "chapter_4_4_significant_digits",
    }

    defaults = _build_default_report_sections(snapshot)
    merged_sections: dict[str, str] = {}
    for key in REPORT_EDITOR_EDITABLE_KEYS:
        raw_value = report_sections.get(key)
        if raw_value is None and key in compatibility_map:
            raw_value = report_sections.get(compatibility_map[key])
        cleaned = str(raw_value).strip() if isinstance(raw_value, str) else ""
        merged_sections[key] = cleaned or defaults.get(key, "")

    normalized["report_sections"] = merged_sections
    return normalized


def _load_latest_draft_payload(account_id: Any) -> dict[str, Any] | None:
    outdir = Path("uploads") / "ai_reports"
    if not outdir.exists():
        return None

    pattern = f"draft_report_{account_id}_*.json"
    draft_files = sorted(outdir.glob(pattern), key=lambda path: path.stat().st_mtime, reverse=True)
    for draft_file in draft_files[:5]:
        try:
            payload = json.loads(draft_file.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            continue
        if isinstance(payload, dict):
            return payload
    return None


def _report_file_prefixes(account_id: Any) -> tuple[str, ...]:
    return (
        f"manual_report_{account_id}_",
        f"draft_report_{account_id}_",
    )


def _list_recent_report_files(account_id: Any, limit: int = 20) -> list[dict[str, Any]]:
    outdir = Path("uploads") / "ai_reports"
    if not outdir.exists():
        return []

    allowed_prefixes = _report_file_prefixes(account_id)
    files: list[dict[str, Any]] = []
    for path in sorted(outdir.iterdir(), key=lambda item: item.stat().st_mtime, reverse=True):
        if not path.is_file():
            continue
        if not path.name.startswith(allowed_prefixes):
            continue
        files.append(
            {
                "name": path.name,
                "size": path.stat().st_size,
                "modified_at": datetime.fromtimestamp(path.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
            }
        )
        if len(files) >= limit:
            break
    return files


def _extract_path_tokens(path: str) -> list[str | int]:
    parts: list[str | int] = []
    for segment in path.split("."):
        segment = segment.strip()
        if not segment:
            continue

        matches = re.finditer(r"([^\[\]]+)|\[(\d+)\]", segment)
        for match in matches:
            key_part = match.group(1)
            index_part = match.group(2)
            if key_part is not None:
                parts.append(key_part)
            elif index_part is not None:
                parts.append(int(index_part))
    return parts


def _resolve_template_value(source: dict, path: str):
    current = source
    for token in _extract_path_tokens(path):
        if isinstance(token, int):
            if not isinstance(current, list):
                return None
            if token < 0 or token >= len(current):
                return None
            current = current[token]
            continue

        if not isinstance(current, dict):
            return None
        if token not in current:
            return None
        current = current[token]
    return current


def _render_template_with_snapshot(template_text: str, snapshot: dict) -> tuple[str, list[str]]:
    missing_tags: list[str] = []

    def _replace(match: re.Match[str]) -> str:
        expr = match.group(1).strip()
        value = _resolve_template_value(snapshot, expr)
        if value is None:
            missing_tags.append(expr)
            return ""
        if isinstance(value, (dict, list)):
            return json.dumps(value, ensure_ascii=False)
        return str(value)

    rendered = TEMPLATE_TAG_PATTERN.sub(_replace, template_text)
    return rendered, sorted(set(missing_tags))


def _collect_scalar_paths(data, prefix: str = "", max_list_items: int = 3) -> list[str]:
    paths: list[str] = []
    if isinstance(data, dict):
        for key, value in data.items():
            next_prefix = f"{prefix}.{key}" if prefix else str(key)
            paths.extend(_collect_scalar_paths(value, next_prefix, max_list_items=max_list_items))
        return paths

    if isinstance(data, list):
        limit = min(len(data), max_list_items)
        for i in range(limit):
            next_prefix = f"{prefix}[{i}]"
            paths.extend(_collect_scalar_paths(data[i], next_prefix, max_list_items=max_list_items))
        return paths

    if prefix:
        paths.append(prefix)
    return paths


def _build_section_34_text(snapshot: dict) -> str:
    devices = snapshot.get("devices_for_section", [])
    if not devices:
        return "資料缺口：尚未建立排放源資料，無法產生 3.4 節內容。"

    lines: list[str] = []
    for d in devices:
        name = d.get("name", "未命名")
        etype = d.get("emission_type", "未分類")
        fuel = d.get("factor_ref_code", "")
        scope = "範疇一" if d.get("scope") == "scope1" else "範疇二"
        lines.append(f"- {name}（排放類型：{etype}，燃料代碼：{fuel}，{scope}）")

    return "\n".join(lines) if lines else "資料缺口：尚未建立排放源資料。"


def _build_markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    if not rows:
        return ""
    header_line = "| " + " | ".join(headers) + " |"
    separator_line = "| " + " | ".join(["---"] * len(headers)) + " |"
    body_lines = ["| " + " | ".join(row) + " |" for row in rows]
    return "\n".join([header_line, separator_line, *body_lines])


def _build_scope_device_table(snapshot: dict[str, Any], scope_key: str) -> str:
    devices = snapshot.get("devices_for_section", [])
    if not isinstance(devices, list):
        return ""

    rows: list[list[str]] = []
    for device in devices:
        if not isinstance(device, dict):
            continue
        if str(device.get("scope") or "") != scope_key:
            continue
        rows.append(
            [
                str(device.get("name") or "未命名設備"),
                str(device.get("emission_type") or "未分類"),
                str(device.get("factor_ref_code") or "-"),
                "範疇一" if scope_key == "scope1" else "範疇二",
            ]
        )
    return _build_markdown_table(["設備名稱", "排放類型", "燃料/來源代碼", "範疇"], rows)


def _build_emission_type_table(snapshot: dict[str, Any]) -> str:
    emission_type_totals = snapshot.get("emission_type_totals", [])
    if not isinstance(emission_type_totals, list):
        return ""
    rows: list[list[str]] = []
    for item in emission_type_totals:
        if not isinstance(item, dict):
            continue
        rows.append([
            str(item.get("emission_type") or "未分類"),
            str(item.get("total_co2e") or 0),
        ])
    return _build_markdown_table(["排放類型", "排放量 (kg CO₂e)"], rows)


def _build_summary_table(snapshot: dict[str, Any]) -> str:
    summary = snapshot.get("summary", {}) if isinstance(snapshot.get("summary"), dict) else {}
    rows = [
        ["盤查年度", str(snapshot.get("inventory_year") or "")],
        ["總排放量 (kg CO₂e)", str(summary.get("total_co2e") or 0)],
        ["排放紀錄筆數", str(summary.get("record_count") or 0)],
        ["設備數量", str(summary.get("device_count") or 0)],
    ]
    return _build_markdown_table(["項目", "內容"], rows)


def _default_report_template() -> str:
    return (
        "# {{ company_summary.company_name }}\n"
        "## {{ inventory_year }} 年溫室氣體盤查報告書\n"
        "\n"
        "盤查期間：{{ inventory_year }} 年 1 月 1 日 ~ {{ inventory_year }} 年 12 月 31 日\n"
        "生成時間：{{ generated_at }}\n"
        "\n"
        "# 第1章 公司簡介與政策聲明\n"
        "## 1.1 前言\n"
        "{{ report_sections.chapter_1_1_preface }}\n"
        "\n"
        "## 1.2 預期用途\n"
        "{{ report_sections.chapter_1_2_expected_use }}\n"
        "\n"
        "## 1.3 公司簡介\n"
        "### 1.3.1 公司基本資料\n"
        "{{ report_sections.chapter_1_3_1_basic_info }}\n"
        "\n"
        "### 1.3.2 企業發展沿革\n"
        "{{ report_sections.chapter_1_3_2_history }}\n"
        "\n"
        "### 1.3.3 主要營業項目\n"
        "{{ report_sections.chapter_1_3_3_main_business }}\n"
        "\n"
        "### 1.3.4 營運概況\n"
        "{{ report_sections.chapter_1_3_4_operations }}\n"
        "\n"
        "# 第2章 盤查邊界設定\n"
        "## 2.1 組織邊界設定\n"
        "{{ report_sections.chapter_2_1_org_boundary }}\n"
        "\n"
        "## 2.2 營運邊界\n"
        "{{ report_sections.chapter_2_2_operational_boundary }}\n"
        "\n"
        "# 第3章 報告溫室氣體排放量\n"
        "## 3.1 溫室氣體排放類型與排放量說明\n"
        "{{ report_sections.chapter_3_1_emission_types }}\n"
        "\n"
        "{{ emission_type_markdown_table }}\n"
        "\n"
        "## 3.2 直接溫室氣體排放（類別1）\n"
        "{{ report_sections.chapter_3_2_category1 }}\n"
        "\n"
        "{{ direct_emission_source_table }}\n"
        "\n"
        "## 3.3 能源間接溫室氣體排放（類別2）\n"
        "{{ report_sections.chapter_3_3_category2 }}\n"
        "\n"
        "{{ indirect_emission_source_table }}\n"
        "\n"
        "## 3.4 溫室氣體總排放量\n"
        "{{ report_sections.chapter_3_4_total }}\n"
        "\n"
        "{{ emission_summary_markdown_table }}\n"
        "\n"
        "# 第4章 數據品質管理\n"
        "## 4.1 量化方法\n"
        "### 4.1.1 活動數據蒐集與轉換方式\n"
        "{{ report_sections.chapter_4_1_1_activity_collection }}\n"
        "\n"
        "### 4.1.2 排放係數與 GWP 換算依據\n"
        "{{ report_sections.chapter_4_1_2_conversion_and_gwp }}\n"
        "\n"
        "## 4.2 量化方法變更說明\n"
        "{{ report_sections.chapter_4_2_method_change }}\n"
        "\n"
        "## 4.3 排放係數與變更說明\n"
        "{{ report_sections.chapter_4_3_factor_change }}\n"
        "\n"
        "## 4.4 有效位數\n"
        "### 4.4.1 活動數據\n"
        "{{ report_sections.chapter_4_4_1_activity_data }}\n"
        "\n"
        "### 4.4.2 排放係數\n"
        "{{ report_sections.chapter_4_4_2_emission_factors }}\n"
        "\n"
        "### 4.4.3 低位熱值\n"
        "{{ report_sections.chapter_4_4_3_lhv }}\n"
        "\n"
        "### 4.4.4 每單位體積或重量之溫室氣體排放量\n"
        "{{ report_sections.chapter_4_4_4_unit_emission }}\n"
        "\n"
        "### 4.4.5 單一排放源之各溫室氣體排放量\n"
        "{{ report_sections.chapter_4_4_5_single_gas }}\n"
        "\n"
        "### 4.4.6 單一排放源之單一溫室氣體排放當量\n"
        "{{ report_sections.chapter_4_4_6_single_co2e }}\n"
        "\n"
        "### 4.4.7 單一排放源之總排放當量\n"
        "{{ report_sections.chapter_4_4_7_source_total }}\n"
        "\n"
        "### 4.4.8 總排放當量彙總\n"
        "{{ report_sections.chapter_4_4_8_grand_total }}\n"
        "\n"
        "## 4.5 重大排放源之資訊流\n"
        "{{ report_sections.chapter_4_5_major_source_flow }}\n"
        "\n"
        "# 第5章 基準年\n"
        "## 5.1 基準年設定\n"
        "{{ report_sections.chapter_5_1_base_year }}\n"
        "\n"
        "{{ base_year_markdown_table }}\n"
        "\n"
        "# 第6章 參考文獻\n"
        "## 6.1 參考文獻\n"
        "{{ report_sections.chapter_6_references }}\n"
        "\n"
        "# 附件\n"
        "## 附件一 公司組織圖\n"
        "{{ report_sections.appendix_1_org_chart }}\n"
        "\n"
        "## 附件二 公司地理位置\n"
        "{{ report_sections.appendix_2_location }}\n"
        "\n"
        "## 附件三 公司廠區配置圖\n"
        "{{ report_sections.appendix_3_layout }}\n"
        "\n"
        "## 附件四 重大排放源之資訊流\n"
        "{{ report_sections.appendix_4_info_flow }}\n"
        "\n"
        "## 附件五 溫室氣體盤查工具\n"
        "{{ report_sections.appendix_5_tools }}\n"
    )


def _build_template_snapshot(snapshot: dict) -> dict:
    template_snapshot = dict(snapshot)
    devices = snapshot.get("devices_for_section", [])
    template_snapshot["devices_for_section"] = devices if isinstance(devices, list) else []
    template_snapshot["section_3_4_text"] = _build_section_34_text(template_snapshot)
    template_snapshot["direct_emission_source_table"] = _build_scope_device_table(template_snapshot, "scope1")
    template_snapshot["indirect_emission_source_table"] = _build_scope_device_table(template_snapshot, "scope2")
    template_snapshot["emission_type_markdown_table"] = _build_emission_type_table(template_snapshot)
    template_snapshot["emission_summary_markdown_table"] = _build_summary_table(template_snapshot)
    template_snapshot["base_year_markdown_table"] = _build_summary_table(template_snapshot)
    return template_snapshot


def _markdown_table_to_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for raw_line in lines:
        stripped = raw_line.strip()
        if not stripped.startswith("|") or not stripped.endswith("|"):
            continue
        rows.append([cell.strip() for cell in stripped.strip("|").split("|")])
    if len(rows) >= 2:
        separator_candidate = rows[1]
        if all(cell.replace(":", "").replace("-", "").strip() == "" for cell in separator_candidate):
            rows = [rows[0], *rows[2:]]
    return rows


def _build_docx_from_rendered_template(rendered_text: str, snapshot: dict[str, Any], output_path: Path) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx 套件，請執行: pip install python-docx==1.1.2") from exc

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft JhengHei"
    normal_style.font.size = Pt(11)
    try:
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    except Exception:
        pass

    company_name = str(snapshot.get("company_summary", {}).get("company_name") or "溫室氣體盤查報告")
    inventory_year = snapshot.get("inventory_year") or ""

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(company_name)
    title_run.bold = True
    title_run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"{inventory_year} 年溫室氣體盤查報告書")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(18)

    period = document.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period.add_run(f"盤查期間：{inventory_year} 年 1 月 1 日 ~ {inventory_year} 年 12 月 31 日")

    document.add_page_break()
    document.add_heading("目錄", level=1)
    for item in REPORT_EDITOR_OUTLINE:
        paragraph = document.add_paragraph()
        if item.get("level") == 2:
            paragraph.paragraph_format.left_indent = Pt(14)
        elif item.get("level") == 3:
            paragraph.paragraph_format.left_indent = Pt(28)
        paragraph.add_run(str(item.get("title") or ""))

    document.add_page_break()

    table_lines: list[str] = []

    def flush_table() -> None:
        nonlocal table_lines
        if not table_lines:
            return
        rows = _markdown_table_to_rows(table_lines)
        table_lines = []
        if not rows:
            return
        max_cols = max(len(row) for row in rows)
        table = document.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"
        for row_index, row in enumerate(rows):
            for col_index in range(max_cols):
                cell_text = row[col_index] if col_index < len(row) else ""
                table.rows[row_index].cells[col_index].text = cell_text
                if row_index == 0:
                    for run in table.rows[row_index].cells[col_index].paragraphs[0].runs:
                        run.bold = True

    for raw_line in rendered_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines.append(stripped)
            continue

        flush_table()

        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
            continue
        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
            continue
        if re.match(r"^\d+\.\s", stripped):
            document.add_paragraph(stripped, style="List Number")
            continue
        document.add_paragraph(stripped)

    flush_table()
    document.save(str(output_path))
    return output_path

def get_session():
    with Session(engine) as session:
        yield session


def _get_user_id(request: Request):
    """從 session 取得當前使用者 ID"""
    return request.session.get("user")


@router.get("/", response_class=HTMLResponse)
async def calculation_page(request: Request, session: Session = Depends(get_session)):
    user_id = _get_user_id(request)
    if not user_id:
        return RedirectResponse(url="/login", status_code=303)

    devices = session.exec(
        select(Device).where(Device.account_id == user_id)
    ).all()
    user_device_ids = [d.id for d in devices]

    records = []
    if user_device_ids:
        records = session.exec(
            select(EmissionRecord)
            .where(EmissionRecord.device_id.in_(user_device_ids))
            .order_by(col(EmissionRecord.record_date).desc())
        ).all()

    device_map = {d.id: d.name for d in devices}
    device_unit_map = {d.id: d.unit for d in devices}
    device_to_code = {d.id: d.factor_ref_code for d in devices}
    device_emission_type_map = {d.id: d.emission_type for d in devices}
    device_by_id = {d.id: d for d in devices}
    device_code_map = {d.id: _derive_device_code_display(d) for d in devices}

    gwp_refs = session.exec(
        select(GWPReference).order_by(GWPReference.gas_name_zh)
    ).all()
    gwp_lookup: dict[str, dict] = {}
    for g in gwp_refs:
        gwp_lookup[g.formula] = {
            "name": g.gas_name_zh,
            "gwp": float(g.gwp_value or 0),
        }

    # --- 直接從 EmissionRecord 與 EmissionFactor604 組合計算過程 ---
    record_calc_map: dict[int, dict] = {}
    for r in records:
        d = device_by_id.get(r.device_id)
        calc: dict = {
            "co2": float(r.co2 or 0),
            "ch4": float(r.ch4 or 0),
            "n2o": float(r.n2o or 0),
            "co2e": float(r.total_co2e or 0),
            "activity_unit": r.activity_unit or r.unit or (d.unit if d else ""),
            "factor_year": r.factor_year,
            "factor_source": r.factor_source,
            "target_year": r.target_year,
        }

        if not d:
            record_calc_map[r.id] = calc
            continue

        etype = (d.emission_type or "").strip()

        if etype == "逸散排放" and d.refrigerant_code:
            gwp_info = gwp_lookup.get(d.refrigerant_code)
            if not gwp_info:
                from services.emission_calculator import _lookup_gwp
                gwp_val = _lookup_gwp(session, d.refrigerant_code)
                gwp_info = {"name": d.refrigerant_code, "gwp": gwp_val}
            rate = get_rate_by_code(d.equipment_category or "") or 0.0
            from services.emission_calculator import _convert_to_kg
            raw_activity = float(r.activity_data or 0)
            activity_unit = r.activity_unit or d.unit or ""
            try:
                activity_kg = _convert_to_kg(raw_activity, activity_unit)
            except ValueError:
                activity_kg = raw_activity
            calc["type"] = "refrigerant"
            calc["refrigerant_name"] = gwp_info.get("name", d.refrigerant_code)
            calc["gwp_value"] = gwp_info.get("gwp", 0)
            calc["emission_rate"] = rate
            calc["activity_value"] = raw_activity
            calc["activity_unit"] = activity_unit
            calc["activity_kg"] = activity_kg
        elif etype == "能源間接排放" or d.factor_ref_code == "ELECTRICITY":
            year = r.factor_year
            factor = session.exec(
                select(EmissionFactor604).where(
                    EmissionFactor604.original_code == "ELECTRICITY",
                    EmissionFactor604.gas_type == "CO2e",
                    EmissionFactor604.year == year,
                )
            ).first()
            if not factor:
                factor = session.exec(
                    select(EmissionFactor604)
                    .where(
                        EmissionFactor604.original_code == "ELECTRICITY",
                        EmissionFactor604.gas_type == "CO2e",
                    )
                    .order_by(EmissionFactor604.year.desc())
                ).first()
            calc["type"] = "electricity"
            calc["factor_value"] = factor.factor_value if factor else 0.0
            calc["factor_unit"] = factor.unit if factor else ""
        else:
            factors = session.exec(
                select(EmissionFactor604).where(
                    EmissionFactor604.original_code == d.factor_ref_code,
                    EmissionFactor604.emission_type == d.emission_type,
                )
            ).all()
            gas_factors = {}
            for f in factors:
                if f.gas_type in ("CO2", "CH4", "N2O"):
                    gas_factors[f.gas_type] = {
                        "value": f.factor_value,
                        "unit": f.unit,
                    }
            calc["type"] = "combustion"
            calc["gas_factors"] = gas_factors

        record_calc_map[r.id] = calc

    # 保持模板相容用的空結構
    factor_detail_map: dict[int, list] = {}
    device_calc_info: dict[int, dict] = {}
    target_year = datetime.now().year

    return templates.TemplateResponse(
        "calculation.html",
        {
            "request": request,
            "records": records,
            "target_year": target_year,
            "device_map": device_map,
            "device_unit_map": device_unit_map,
            "device_to_code": device_to_code,
            "device_emission_type_map": device_emission_type_map,
            "device_code_map": device_code_map,
            "factor_detail_map": factor_detail_map,
            "device_calc_info": device_calc_info,
            "record_calc_map": record_calc_map,
        },
    )


@router.post("/delete/{record_id}")
async def delete_record(
    request: Request, record_id: int, session: Session = Depends(get_session)
):
    user_id = _get_user_id(request)
    if not user_id:
        return RedirectResponse(url="/login", status_code=303)

    record = session.get(EmissionRecord, record_id)
    if record:
        device = session.get(Device, record.device_id)
        if device and device.account_id == user_id:
            add_change_log(
                session=session,
                module="calculation",
                entity_name="EmissionRecord",
                record_key=str(record.id),
                action_type="DELETE",
                changed_by=str(user_id),
                change_details=(
                    f"device_id={record.device_id}, record_date={record.record_date}, "
                    f"activity_data={record.activity_data}, total_co2e={record.total_co2e}"
                ),
            )
            session.delete(record)
            session.commit()
    return RedirectResponse(url="/calculation/", status_code=303)


@router.get("/logs")
async def get_logs(
    request: Request,
    limit: int = 50,
    module: str = "all",
    session: Session = Depends(get_session),
):
    user_id = _get_user_id(request)
    if not user_id:
        return JSONResponse(status_code=403, content={"error": "未登入"})

    normalized_limit = min(max(limit, 1), 200)
    query = select(DataChangeLog).where(DataChangeLog.changed_by == str(user_id))
    if module != "all":
        query = query.where(DataChangeLog.module == module)
    logs = session.exec(
        query.order_by(col(DataChangeLog.changed_at).desc()).limit(normalized_limit)
    ).all()
    payload = [
        {
            "id": log.id,
            "module": log.module,
            "entity_name": log.entity_name,
            "record_key": log.record_key,
            "action_type": log.action_type,
            "changed_by": log.changed_by,
            "changed_at": log.changed_at.isoformat(sep=" ", timespec="seconds"),
            "change_details": log.change_details,
        }
        for log in logs
    ]
    return JSONResponse(payload)


@router.get("/report", response_class=HTMLResponse)
async def report_edit_page(request: Request, session: Session = Depends(get_session)):
    # import result module dynamically to avoid circular imports
    from routers import result as result_mod

    account_id = request.session.get("user")
    if not account_id:
        return RedirectResponse(url="/login", status_code=303)
    snapshot = result_mod._build_report_snapshot(session, account_id=account_id)

    draft_payload = _load_latest_draft_payload(account_id)
    template_snapshot = _build_template_snapshot(snapshot)
    try:
        initial_payload = result_mod._normalize_report_payload(draft_payload or {}, snapshot)
    except Exception:
        initial_payload = draft_payload or {}
    initial_payload = _ensure_report_sections(initial_payload, template_snapshot)
    template_snapshot["report_sections"] = initial_payload.get("report_sections", {})

    payload_json = json.dumps(initial_payload, ensure_ascii=False, indent=2)
    snapshot_json = json.dumps(snapshot, ensure_ascii=False, indent=2)
    template_text = _default_report_template()
    rendered_template, missing_template_tags = _render_template_with_snapshot(template_text, template_snapshot)

    return templates.TemplateResponse(
        "report_edit.html",
        {
            "request": request,
            "snapshot": snapshot,
            "snapshot_json": snapshot_json,
            "payload_json": payload_json,
            "template_text": template_text,
            "rendered_template": rendered_template,
            "missing_template_tags": missing_template_tags,
            "template_paths": _collect_scalar_paths(template_snapshot),
            "report_outline": REPORT_EDITOR_OUTLINE,
            "report_files": _list_recent_report_files(account_id),
        },
    )


@router.post("/report/render", response_class=HTMLResponse)
async def render_report_template(request: Request, session: Session = Depends(get_session)):
    from routers import result as result_mod

    account_id = request.session.get("user")
    if not account_id:
        return RedirectResponse(url="/login", status_code=303)
    snapshot = result_mod._build_report_snapshot(session, account_id=account_id)
    template_snapshot = _build_template_snapshot(snapshot)

    try:
        initial_payload = result_mod._normalize_report_payload({}, snapshot)
    except Exception:
        initial_payload = {}

    form = await request.form()
    template_text = str(form.get("template_text") or _default_report_template())
    payload_json = str(form.get("payload_json") or json.dumps(initial_payload, ensure_ascii=False, indent=2))
    try:
        payload = json.loads(payload_json)
    except Exception:
        payload = initial_payload
    payload = _ensure_report_sections(payload, template_snapshot)
    template_snapshot["report_sections"] = payload.get("report_sections", {})
    payload_json = json.dumps(payload, ensure_ascii=False, indent=2)
    snapshot_json = json.dumps(snapshot, ensure_ascii=False, indent=2)

    rendered_template, missing_template_tags = _render_template_with_snapshot(template_text, template_snapshot)

    return templates.TemplateResponse(
        "report_edit.html",
        {
            "request": request,
            "snapshot": snapshot,
            "snapshot_json": snapshot_json,
            "payload_json": payload_json,
            "template_text": template_text,
            "rendered_template": rendered_template,
            "missing_template_tags": missing_template_tags,
            "template_paths": _collect_scalar_paths(template_snapshot),
            "report_outline": REPORT_EDITOR_OUTLINE,
            "report_files": _list_recent_report_files(account_id),
        },
    )


@router.post("/report/save")
async def save_report(request: Request, session: Session = Depends(get_session)):
    from routers import result as result_mod

    form = await request.form()
    payload_json = form.get("payload_json", "")
    template_text = str(form.get("template_text", "") or "")
    account_id = request.session.get("user")
    if not account_id:
        return RedirectResponse(url="/login", status_code=303)
    snapshot = result_mod._build_report_snapshot(session, account_id=account_id)
    snapshot_json = json.dumps(snapshot, ensure_ascii=False, indent=2)
    template_snapshot = _build_template_snapshot(snapshot)
    error = None
    try:
        payload = json.loads(payload_json)
    except Exception as exc:
        error = str(exc)
        rendered_template, missing_template_tags = _render_template_with_snapshot(
            template_text or _default_report_template(),
            template_snapshot,
        )
        return templates.TemplateResponse(
            "report_edit.html",
            {
                "request": request,
                "snapshot": snapshot,
                "snapshot_json": snapshot_json,
                "payload_json": payload_json,
                "error": error,
                "template_text": template_text or _default_report_template(),
                "rendered_template": rendered_template,
                "missing_template_tags": missing_template_tags,
                "template_paths": _collect_scalar_paths(template_snapshot),
                "report_outline": REPORT_EDITOR_OUTLINE,
                "report_files": _list_recent_report_files(account_id),
            },
        )

    payload = _ensure_report_sections(payload, template_snapshot)
    template_snapshot["report_sections"] = payload.get("report_sections", {})

    outdir = Path("uploads") / "ai_reports"
    outdir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = outdir / f"manual_report_{account_id}_{timestamp}.json"
    filename.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    if template_text.strip():
        rendered_template, _ = _render_template_with_snapshot(template_text, template_snapshot)
        rendered_path = outdir / f"manual_report_{account_id}_{timestamp}.md"
        rendered_path.write_text(rendered_template, encoding="utf-8")

    return RedirectResponse(url="/calculation/report", status_code=303)


@router.get("/report/download/{filename}")
async def download_saved_report(filename: str, request: Request):
    account_id = request.session.get("user")
    if not account_id:
        return RedirectResponse(url="/login", status_code=303)

    safe_name = Path(filename).name
    if not safe_name.startswith(_report_file_prefixes(account_id)):
        raise HTTPException(status_code=403, detail="無權限下載此報告檔案")

    file_path = Path("uploads") / "ai_reports" / safe_name
    if not file_path.exists() or not file_path.is_file():
        raise HTTPException(status_code=404, detail="找不到報告檔案")

    media_type = "application/octet-stream"
    if safe_name.endswith(".docx"):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif safe_name.endswith(".json"):
        media_type = "application/json"
    elif safe_name.endswith(".md"):
        media_type = "text/markdown"

    return FileResponse(str(file_path), media_type=media_type, filename=safe_name)


@router.post("/report/export-docx")
async def export_report_docx(request: Request, session: Session = Depends(get_session)):
    from routers import result as result_mod

    account_id = request.session.get("user")
    if not account_id:
        return RedirectResponse(url="/login", status_code=303)

    form = await request.form()
    payload_json = str(form.get("payload_json") or "{}")
    template_text = str(form.get("template_text") or _default_report_template())

    try:
        payload = json.loads(payload_json)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail=f"payload_json 格式錯誤：{exc}") from exc

    snapshot = result_mod._build_report_snapshot(session, account_id=account_id)
    template_snapshot = _build_template_snapshot(snapshot)
    payload = _ensure_report_sections(payload, template_snapshot)
    template_snapshot["report_sections"] = payload.get("report_sections", {})

    rendered_template, _ = _render_template_with_snapshot(template_text, template_snapshot)
    outdir = Path("uploads") / "ai_reports"
    outdir.mkdir(parents=True, exist_ok=True)
    file_path = outdir / f"manual_report_{account_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    _build_docx_from_rendered_template(rendered_template, template_snapshot, file_path)

    download_name = f"ghg_report_{snapshot.get('inventory_year')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return FileResponse(
        str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=download_name,
    )


@router.post("/report/save-draft")
async def save_report_draft(request: Request, session: Session = Depends(get_session)):
    account_id = request.session.get("user")
    if not account_id:
        raise HTTPException(status_code=401, detail="尚未登入")

    body = await request.json()
    payload = body.get("payload") if isinstance(body, dict) else None
    if not isinstance(payload, dict):
        raise HTTPException(status_code=400, detail="payload 格式錯誤，必須為 JSON object")

    outdir = Path("uploads") / "ai_reports"
    outdir.mkdir(parents=True, exist_ok=True)
    filename = outdir / f"draft_report_{account_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    filename.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    return {
        "ok": True,
        "message": "草稿已儲存",
        "file": str(filename),
    }
