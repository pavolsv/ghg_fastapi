import json
import os
from datetime import datetime
from pathlib import Path
from threading import Lock
from typing import Any
from uuid import uuid4

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Request
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse
from fastapi.templating import Jinja2Templates
from sqlmodel import Session, select

from database import engine
from model import CompanyInfo, Device, EmissionRecord, GWPReference, Report
from datetime import datetime

# ...
def _get_target_year(session: Session) -> int:
    from model import EmissionFactor604
    from sqlmodel import select
    all_factors = session.exec(select(EmissionFactor604)).all()
    years = [f.year for f in all_factors if f.year is not None]
    return max(years, default=datetime.now().year)

router = APIRouter(prefix="/result", tags=["result"])
templates = Jinja2Templates(directory="templates")
AI_REPORT_TASKS: dict[str, dict[str, Any]] = {}
AI_REPORT_LOCK = Lock()
AI_REPORT_OUTPUT_DIR = Path("uploads") / "ai_reports"
AI_REPORT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
AI_REPORT_LOCAL_ENV_PATH = Path("env.local.txt")


def get_session():
    with Session(engine) as session:
        yield session


def _set_task(task_id: str, **fields: Any) -> None:
    with AI_REPORT_LOCK:
        current = AI_REPORT_TASKS.get(task_id, {})
        current.update(fields)
        AI_REPORT_TASKS[task_id] = current


def _get_task(task_id: str) -> dict[str, Any] | None:
    with AI_REPORT_LOCK:
        task = AI_REPORT_TASKS.get(task_id)
        return dict(task) if task else None


def _load_local_env_settings() -> dict[str, str]:
    if not AI_REPORT_LOCAL_ENV_PATH.exists():
        return {}

    try:
        raw_text = AI_REPORT_LOCAL_ENV_PATH.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw_text = AI_REPORT_LOCAL_ENV_PATH.read_text(encoding="utf-8-sig")

    settings: dict[str, str] = {}
    for raw_line in raw_text.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#"):
            continue

        if line.lower().startswith("export "):
            line = line[7:].strip()
        if line.startswith("$env:"):
            line = line[5:].strip()

        if "=" not in line:
            continue

        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip()
        if not key:
            continue

        if (value.startswith('"') and value.endswith('"')) or (
            value.startswith("'") and value.endswith("'")
        ):
            value = value[1:-1]

        settings[key] = value

    return settings


def _get_env_value(*keys: str, default: str | None = None) -> str | None:
    for key in keys:
        env_value = os.getenv(key)
        if env_value and env_value.strip():
            return env_value.strip()

    local_settings = _load_local_env_settings()
    for key in keys:
        local_value = local_settings.get(key)
        if local_value and local_value.strip():
            return local_value.strip()

    return default


def _resolve_llm_runtime_config() -> dict[str, str]:
    provider = (_get_env_value("LLM_PROVIDER", default="gemini") or "gemini").strip().lower()

    if provider == "gemini":
        api_key = _get_env_value("GEMINI_API_KEY", "GOOGLE_API_KEY")
        if not api_key:
            raise RuntimeError("未設定 GEMINI_API_KEY (或 GOOGLE_API_KEY)，無法生成 AI 報告。")

        return {
            "provider": provider,
            "api_key": api_key,
            "base_url": _get_env_value(
                "GEMINI_BASE_URL",
                default="https://generativelanguage.googleapis.com/v1beta/openai/",
            )
            or "https://generativelanguage.googleapis.com/v1beta/openai/",
            "model": _get_env_value("GEMINI_MODEL", default="gemini-2.5-flash") or "gemini-2.5-flash",
        }

    if provider == "openai":
        api_key = _get_env_value("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError("未設定 OPENAI_API_KEY，無法生成 AI 報告。")

        return {
            "provider": provider,
            "api_key": api_key,
            "model": _get_env_value("OPENAI_MODEL", default="gpt-4o-mini") or "gpt-4o-mini",
        }

    raise RuntimeError("不支援的 LLM_PROVIDER，請使用 openai 或 gemini。")


def _build_report_snapshot(session: Session, account_id: int | None = None) -> dict[str, Any]:
    records = session.exec(select(EmissionRecord)).all()
    devices = session.exec(select(Device)).all()

    if account_id is None:
        companies = session.exec(select(CompanyInfo)).all()
    else:
        companies = session.exec(
            select(CompanyInfo).where(CompanyInfo.account_id == account_id)
        ).all()

    device_map = {device.id: device for device in devices}

    emission_type_totals_raw: dict[str, float] = {}
    device_totals_raw: dict[int, float] = {}
    scope_source_counts_raw: dict[str, int] = {}
    total_co2e = 0.0

    for record in records:
        device = device_map.get(record.device_id)
        emission_type = (device.emission_type if device else "未分類") or "未分類"
        device_name = (device.name if device else f"設備#{record.device_id}") or "未命名設備"
        emission_type = emission_type.strip() or "未分類"
        device_name = device_name.strip() or "未命名設備"

        co2e_value = float(record.total_co2e or 0.0)
        total_co2e += co2e_value
        emission_type_totals_raw[emission_type] = emission_type_totals_raw.get(emission_type, 0.0) + co2e_value
        device_totals_raw[record.device_id] = device_totals_raw.get(record.device_id, 0.0) + co2e_value

        # scope 統計（來自 Device）
        dev_scope = (device.scope if device else "scope1") or "scope1"
        display_scope = "範疇一" if dev_scope == "scope1" else "範疇二"
        scope_source_counts_raw[display_scope] = scope_source_counts_raw.get(display_scope, 0) + 1

    emission_type_totals = [
        {"emission_type": key, "total_co2e": round(value, 4)}
        for key, value in sorted(
            emission_type_totals_raw.items(),
            key=lambda item: item[1],
            reverse=True,
        )
    ]
    top_devices = [
        {
            "device_name": (device_map.get(did).name if device_map.get(did) else f"設備#{did}"),
            "total_co2e": round(value, 4),
        }
        for did, value in sorted(
            device_totals_raw.items(),
            key=lambda item: item[1],
            reverse=True,
        )[:5]
    ]

    scope_source_counts = [
        {"scope": key, "source_count": value}
        for key, value in sorted(scope_source_counts_raw.items(), key=lambda item: item[0])
    ]

    company = companies[0] if companies else None
    company_summary = {
        "company_name": (company.company_name if company else None),
        "tax_id": (company.tax_id if company else None),
        "address": (company.address if company else None),
        "owner": (company.owner if company else None),
        "contact_person": (company.contact_person if company else None),
        "email": (company.email if company else None),
        "telephone": (company.telephone if company else None),
    }

    numeric_sources: dict[str, float] = {
        "summary.total_co2e": round(total_co2e, 4),
        "summary.record_count": float(len(records)),
        "summary.device_count": float(len(devices)),
    }

    for index, item in enumerate(emission_type_totals):
        numeric_sources[f"emission_type_totals[{index}].total_co2e"] = float(item["total_co2e"])
    for index, item in enumerate(top_devices):
        numeric_sources[f"top_devices[{index}].total_co2e"] = float(item["total_co2e"])
    for index, item in enumerate(scope_source_counts):
        numeric_sources[f"scope_source_counts[{index}].source_count"] = float(item["source_count"])

    return {
        "inventory_year": _get_target_year(session),
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "summary": {
            "total_co2e": round(total_co2e, 4),
            "record_count": len(records),
            "device_count": len(devices),
        },
        "company_summary": company_summary,
        "calculation_summary": {
            "factor_source": "EmissionFactor604",
            "gwp_version": "AR5",
        },
        "scope_source_counts": scope_source_counts,
        "emission_type_totals": emission_type_totals,
        "top_devices": top_devices,
        "numeric_sources": numeric_sources,
    }


def _normalize_report_payload(payload: dict[str, Any], snapshot: dict[str, Any]) -> dict[str, Any]:
    normalized = dict(payload or {})
    compliance_sections = normalized.get("compliance_template_sections")
    if not isinstance(compliance_sections, dict):
        compliance_sections = {}

    summary = snapshot.get("summary", {}) if isinstance(snapshot.get("summary"), dict) else {}
    company_summary = snapshot.get("company_summary", {}) if isinstance(snapshot.get("company_summary"), dict) else {}
    boundary_summary = snapshot.get("boundary_summary", {}) if isinstance(snapshot.get("boundary_summary"), dict) else {}
    calculation_summary = (
        snapshot.get("calculation_summary", {})
        if isinstance(snapshot.get("calculation_summary"), dict)
        else {}
    )
    scope_source_counts = snapshot.get("scope_source_counts", [])
    if not isinstance(scope_source_counts, list):
        scope_source_counts = []
    emission_type_totals = snapshot.get("emission_type_totals", [])
    if not isinstance(emission_type_totals, list):
        emission_type_totals = []
    reduction_actions = normalized.get("reduction_actions")
    if not isinstance(reduction_actions, list):
        reduction_actions = []

    def _clean_string_list(raw: Any) -> list[str]:
        if not isinstance(raw, list):
            return []
        values: list[str] = []
        for item in raw:
            text = str(item).strip()
            if text:
                values.append(text)
        return values

    def _fallback_company_basic_info() -> list[str]:
        lines: list[str] = []
        company_name = company_summary.get("company_name")
        address = company_summary.get("address")
        owner = company_summary.get("owner")
        if company_name:
            lines.append(f"事業名稱：{company_name}")
        if address:
            lines.append(f"事業地址：{address}")
        if owner:
            lines.append(f"事業負責人：{owner}")
        tax_id = company_summary.get("tax_id")
        if tax_id:
            lines.append(f"統一編號：{tax_id}")
        if not lines:
            lines.append("資料缺口：未提供事業名稱、地址與負責人等公司基本資料。")
        return lines

    def _fallback_boundary_setting() -> list[str]:
        lines: list[str] = []
        boundary_count = int(boundary_summary.get("boundary_count", 0) or 0)
        source_count = int(boundary_summary.get("source_count", 0) or 0)
        if boundary_count > 0:
            lines.append(f"盤查邊界數量：{boundary_count}")
        else:
            lines.append("資料缺口：未建立廠（場）或組織邊界資料。")

        if source_count > 0:
            lines.append(f"納入邊界之排放源數量：{source_count}")

        if scope_source_counts:
            scope_parts: list[str] = []
            for item in scope_source_counts:
                if not isinstance(item, dict):
                    continue
                scope_name = str(item.get("scope", "未分類")).strip() or "未分類"
                scope_count = int(item.get("source_count", 0) or 0)
                scope_parts.append(f"{scope_name}{scope_count}項")
            if scope_parts:
                lines.append("範疇分布：" + "、".join(scope_parts))
        else:
            lines.append("資料缺口：缺少範疇一/二排放源分類資料。")
        return lines

    def _fallback_source_identification() -> list[str]:
        lines: list[str] = []
        top_types = emission_type_totals[:3]
        for item in top_types:
            if not isinstance(item, dict):
                continue
            emission_type = str(item.get("emission_type", "未分類")).strip() or "未分類"
            total_value = item.get("total_co2e", 0)
            lines.append(f"排放類型：{emission_type}，排放量：{total_value} kg CO₂e")
        if not lines:
            lines.append("資料缺口：尚無可用之排放源鑑別或排放類型統計資料。")
        return lines

    def _fallback_emission_calculation() -> list[str]:
        lines: list[str] = []
        record_count = int(summary.get("record_count", 0) or 0)
        total_co2e = summary.get("total_co2e", 0)
        if record_count > 0:
            lines.append(f"年排放量總計：{total_co2e} kg CO₂e（依既有紀錄彙總）")
            lines.append(f"排放量計算紀錄筆數：{record_count}")
            lines.append("計算依據：採用環保署公告溫室氣體排放係數表（EmissionFactor604）與 AR5 GWP 值。")
        else:
            lines.append("資料缺口：尚無排放量計算紀錄，無法產出年排放量。")
        return lines

    def _fallback_other_regulatory_items() -> list[str]:
        lines: list[str] = []
        for action in reduction_actions[:3]:
            if not isinstance(action, dict):
                continue
            title = str(action.get("title", "")).strip()
            detail = str(action.get("detail", "")).strip()
            if title and detail:
                lines.append(f"減量措施：{title}。說明：{detail}")
            elif title:
                lines.append(f"減量措施：{title}")
        if not lines:
            lines.append("資料缺口：尚未提供事業執行減量措施及佐證說明。")
        return lines

    def _ensure_section(key: str, fallback: list[str]) -> None:
        cleaned = _clean_string_list(compliance_sections.get(key))
        compliance_sections[key] = cleaned if cleaned else fallback

    _ensure_section("company_basic_info", _fallback_company_basic_info())
    _ensure_section("boundary_setting", _fallback_boundary_setting())
    _ensure_section("source_identification", _fallback_source_identification())
    _ensure_section("emission_calculation", _fallback_emission_calculation())
    _ensure_section("other_regulatory_items", _fallback_other_regulatory_items())

    normalized["compliance_template_sections"] = compliance_sections
    return normalized


def _generate_report_payload(snapshot: dict[str, Any]) -> dict[str, Any]:
    try:
        from openai import OpenAI
    except ImportError as exc:
        raise RuntimeError("缺少 openai 套件，請執行: pip install openai") from exc

    config = _resolve_llm_runtime_config()
    provider = config["provider"]
    if provider == "gemini":

        # Gemini 提供 OpenAI-compatible endpoint，可沿用既有訊息格式與 JSON 模式。
        client = OpenAI(
            api_key=config["api_key"],
            base_url=config["base_url"],
        )
        model = config["model"]
    elif provider == "openai":
        client = OpenAI(api_key=config["api_key"])
        model = config["model"]
    else:
        raise RuntimeError("不支援的 LLM_PROVIDER，請使用 openai 或 gemini。")

    system_prompt = (
        "你是企業溫室氣體盤查顧問。"
        "你只能使用使用者提供的 JSON 事實資料。"
        "嚴禁捏造數字、法規條文、額外排放係數。"
        "若資料不足，必須明確寫出資料缺口。"
        "你必須輸出單一 JSON object，不能有 markdown code fence。"
        "輸出欄位固定為：executive_summary(字串),"
        "compliance_template_sections(物件，且必含 company_basic_info/boundary_setting/source_identification/emission_calculation/other_regulatory_items 五個字串陣列),"
        "current_state_findings(字串陣列),"
        "reduction_actions(物件陣列，每項含 title/detail/expected_impact),"
        "assumptions(字串陣列),"
        "data_gaps(字串陣列),"
        "citations(物件陣列，每項含 metric/source_path/value)。"
        "請盡量貼合盤查報告書章節範本，內容順序依序為：公司基本資料、盤查邊界設定、排放源鑑別、排放量計算、其他主管機關規定事項。"
        "若任何章節資料不足，該章節至少要有一行以『資料缺口：』開頭的說明。"
        "citations.source_path 必須對應到輸入中的 numeric_sources 鍵名，"
        "citations.value 必須填對應數值。"
    )
    user_prompt = json.dumps(snapshot, ensure_ascii=False)

    response = client.chat.completions.create(
        model=model,
        temperature=0.2,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )

    content = (response.choices[0].message.content or "").strip()
    if not content:
        raise RuntimeError("LLM 沒有回傳內容。")

    try:
        return json.loads(content)
    except json.JSONDecodeError as exc:
        raise RuntimeError("LLM 回傳格式不是有效 JSON。") from exc


def _validate_report_payload(payload: dict[str, Any], snapshot: dict[str, Any]) -> None:
    if not isinstance(payload, dict):
        raise ValueError("報告內容格式錯誤：payload 必須為 JSON object。")

    required_string_fields = ["executive_summary"]
    required_list_fields = [
        "current_state_findings",
        "reduction_actions",
        "assumptions",
        "data_gaps",
        "citations",
    ]

    for field in required_string_fields:
        value = payload.get(field)
        if not isinstance(value, str) or not value.strip():
            raise ValueError(f"報告欄位缺失或格式錯誤：{field}")

    for field in required_list_fields:
        value = payload.get(field)
        if not isinstance(value, list):
            raise ValueError(f"報告欄位缺失或格式錯誤：{field}")

    compliance_sections = payload.get("compliance_template_sections")
    if not isinstance(compliance_sections, dict):
        raise ValueError("報告欄位缺失或格式錯誤：compliance_template_sections")

    required_compliance_sections = [
        "company_basic_info",
        "boundary_setting",
        "source_identification",
        "emission_calculation",
        "other_regulatory_items",
    ]
    for section in required_compliance_sections:
        section_value = compliance_sections.get(section)
        if not isinstance(section_value, list):
            raise ValueError(f"盤查章節欄位缺失或格式錯誤：compliance_template_sections.{section}")

    numeric_sources = snapshot.get("numeric_sources", {})
    citations = payload.get("citations", [])
    if not citations:
        raise ValueError("報告缺少 citations，無法驗證數值來源。")

    matched_citations = 0
    for index, citation in enumerate(citations):
        if not isinstance(citation, dict):
            raise ValueError(f"citations[{index}] 格式錯誤")

        source_path = citation.get("source_path")
        value = citation.get("value")
        if not isinstance(source_path, str) or source_path not in numeric_sources:
            raise ValueError(f"citations[{index}] source_path 不存在於 numeric_sources")
        if not isinstance(value, (int, float, str)):
            raise ValueError(f"citations[{index}] value 不是有效數字")

        try:
            numeric_value = float(value)
        except (TypeError, ValueError) as exc:
            raise ValueError(f"citations[{index}] value 不是有效數字") from exc

        source_value = float(numeric_sources[source_path])
        if abs(source_value - numeric_value) > 1e-3:
            raise ValueError(
                f"citations[{index}] value 與資料來源不一致 (source={source_value}, output={numeric_value})"
            )
        matched_citations += 1

    if matched_citations == 0:
        raise ValueError("報告沒有可驗證的數值引用。")


def _build_docx_report(payload: dict[str, Any], snapshot: dict[str, Any], task_id: str) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx 套件，請執行: pip install python-docx==1.1.2") from exc

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft JhengHei"
    normal_style.font.size = Pt(11)
    try:
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    except Exception:
        pass

    company_name = str(snapshot.get("company_summary", {}).get("company_name") or "溫室氣體盤查報告")
    inventory_year = snapshot.get("inventory_year")
    total_co2e = snapshot.get("summary", {}).get("total_co2e")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(company_name)
    title_run.bold = True
    title_run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"{inventory_year} 年溫室氣體盤查報告書")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(18)

    period = document.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period.add_run(f"盤查期間：{inventory_year} 年 1 月 1 日 ~ {inventory_year} 年 12 月 31 日")

    document.add_page_break()
    document.add_heading("目錄", level=1)
    toc_lines = [
        "第1章 公司基本資料",
        "  1.1 公司基本資料",
        "第2章 盤查邊界設定",
        "  2.1 邊界設定",
        "第3章 排放源與排放量",
        "  3.1 排放源鑑別",
        "  3.2 排放量計算",
        "  3.3 溫室氣體總排放量",
        "第4章 其他規定事項與管理建議",
        "  4.1 其他主管機關規定事項",
        "  4.2 現況洞察",
        "  4.3 減量建議與策略",
        "第5章 假設條件與資料缺口",
        "第6章 數值引用",
    ]
    for line in toc_lines:
        document.add_paragraph(line)

    document.add_page_break()
    document.add_heading("第1章 公司基本資料", level=1)
    document.add_heading("1.1 執行摘要", level=2)
    document.add_paragraph(str(payload.get("executive_summary", "")))

    compliance_sections = payload.get("compliance_template_sections", {})

    document.add_heading("1.2 公司基本資料", level=2)
    for item in compliance_sections.get("company_basic_info", []):
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("第2章 盤查邊界設定", level=1)
    document.add_heading("2.1 邊界設定", level=2)
    for item in compliance_sections.get("boundary_setting", []):
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("第3章 排放源與排放量", level=1)
    document.add_heading("3.1 排放源鑑別", level=2)
    for item in compliance_sections.get("source_identification", []):
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("3.2 排放量計算", level=2)
    for item in compliance_sections.get("emission_calculation", []):
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("3.3 溫室氣體總排放量", level=2)
    document.add_paragraph(f"本公司 {inventory_year} 年溫室氣體總排放量為 {total_co2e} kg CO₂e。")

    document.add_heading("第4章 其他規定事項與管理建議", level=1)
    document.add_heading("4.1 其他主管機關規定事項", level=2)
    for item in compliance_sections.get("other_regulatory_items", []):
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("4.2 現況分析與洞察", level=2)
    for item in payload.get("current_state_findings", []):
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("4.3 減碳建議與策略", level=2)
    for item in payload.get("reduction_actions", []):
        if isinstance(item, dict):
            title = str(item.get("title", "未命名建議"))
            detail = str(item.get("detail", ""))
            impact = str(item.get("expected_impact", ""))
            document.add_paragraph(f"{title}", style="List Number")
            if detail:
                document.add_paragraph(f"說明：{detail}")
            if impact:
                document.add_paragraph(f"預期效益：{impact}")
        else:
            document.add_paragraph(str(item), style="List Number")

    document.add_heading("第5章 假設條件與資料缺口", level=1)
    document.add_heading("5.1 假設條件", level=2)
    for item in payload.get("assumptions", []):
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("5.2 資料缺口", level=2)
    for item in payload.get("data_gaps", []):
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("第6章 數值引用", level=1)
    for citation in payload.get("citations", []):
        metric = citation.get("metric", "") if isinstance(citation, dict) else ""
        source_path = citation.get("source_path", "") if isinstance(citation, dict) else ""
        value = citation.get("value", "") if isinstance(citation, dict) else ""
        document.add_paragraph(f"{metric}: {value} (來源: {source_path})", style="List Bullet")

    file_path = AI_REPORT_OUTPUT_DIR / f"ai_report_{task_id}.docx"
    document.save(str(file_path))
    return file_path


def _run_ai_report_task(task_id: str, snapshot: dict[str, Any]) -> None:
    _set_task(task_id, status="processing", message="正在呼叫 LLM 生成報告...")
    try:
        payload = _generate_report_payload(snapshot)
        payload = _normalize_report_payload(payload, snapshot)
        _set_task(task_id, status="processing", message="正在驗證報告內容...")
        _validate_report_payload(payload, snapshot)

        _set_task(task_id, status="processing", message="正在轉換為 Word 檔案...")
        file_path = _build_docx_report(payload, snapshot, task_id)
        _set_task(
            task_id,
            status="completed",
            message="報告生成完成，可下載。",
            completed_at=datetime.now().isoformat(),
            file_path=str(file_path),
        )
    except Exception as exc:
        _set_task(
            task_id,
            status="failed",
            message="報告生成失敗，請檢查設定後再試。",
            completed_at=datetime.now().isoformat(),
            error=str(exc),
        )


@router.get("/", response_class=HTMLResponse)
async def result_page(request: Request, session: Session = Depends(get_session)):
    records = session.exec(select(EmissionRecord)).all()
    devices = session.exec(select(Device)).all()
    device_map = {device.id: device for device in devices}

    # --- 直接從 EmissionRecord 彙總（不再即時重算） ---
    scope_order = ["scope1", "scope2"]
    scope_display = {
        "scope1": "範疇一：直接排放",
        "scope2": "範疇二：能源間接排放",
    }
    emission_type_to_scope: dict[str, str] = {
        "固定燃燒": "scope1",
        "移動燃燒": "scope1",
        "逸散排放": "scope1",
        "能源間接排放": "scope2",
    }

    scope_data: list[dict] = []
    scope_totals: dict[str, float] = {"scope1": 0.0, "scope2": 0.0}
    emission_type_totals_raw: dict[str, float] = {}

    for scope_key in scope_order:
        emission_types_in_scope: dict[str, dict] = {}

        for record in records:
            device = device_map.get(record.device_id)
            etype = ((device.emission_type if device else "未分類") or "未分類").strip()
            dev_name = ((device.name if device else f"設備#{record.device_id}") or "未命名").strip()
            device_key = record.device_id
            factor_code = (device.factor_ref_code if device else "") or ""

            if emission_type_to_scope.get(etype, "scope1") != scope_key:
                continue

            co2e_val = float(record.total_co2e or 0.0)
            gases = {
                "CO2": float(record.co2 or 0),
                "CH4": float(record.ch4 or 0),
                "N2O": float(record.n2o or 0),
                "CO2e": co2e_val,
            }

            emission_type_totals_raw[etype] = emission_type_totals_raw.get(etype, 0.0) + co2e_val

            if etype not in emission_types_in_scope:
                emission_types_in_scope[etype] = {
                    "emission_type": etype,
                    "devices": {},
                    "type_total_co2e": 0.0,
                    "type_gas_totals": {"CO2": 0.0, "CH4": 0.0, "N2O": 0.0},
                }

            et_data = emission_types_in_scope[etype]
            et_data["type_total_co2e"] += co2e_val
            for g in ("CO2", "CH4", "N2O"):
                et_data["type_gas_totals"][g] = et_data["type_gas_totals"].get(g, 0.0) + gases.get(g, 0.0)

            if device_key not in et_data["devices"]:
                et_data["devices"][device_key] = {
                    "device_name": dev_name,
                    "factor_code": factor_code,
                    "records": [],
                    "total_co2e": 0.0,
                    "gas_totals": {"CO2": 0.0, "CH4": 0.0, "N2O": 0.0},
                }

            dev_data = et_data["devices"][device_key]
            dev_data["records"].append({
                "record_date": record.record_date,
                "activity_data": record.activity_data,
                "unit": record.unit or "",
                "target_year": record.target_year,
                "co2e": co2e_val,
                "gases": gases,
            })
            dev_data["total_co2e"] += co2e_val
            for g in ("CO2", "CH4", "N2O"):
                dev_data["gas_totals"][g] = dev_data["gas_totals"].get(g, 0.0) + gases.get(g, 0.0)
            scope_totals[scope_key] += co2e_val

        type_list = []
        for etype_key in sorted(emission_types_in_scope.keys()):
            et = emission_types_in_scope[etype_key]
            device_list = []
            for dev in sorted(et["devices"].values(), key=lambda d: d["device_name"]):
                dev["total_co2e"] = round(dev["total_co2e"], 4)
                for g in ("CO2", "CH4", "N2O"):
                    dev["gas_totals"][g] = round(dev["gas_totals"][g], 4)
                device_list.append(dev)
            for g in ("CO2", "CH4", "N2O"):
                et["type_gas_totals"][g] = round(et["type_gas_totals"][g], 4)
            type_list.append({
                "emission_type": etype_key,
                "devices": device_list,
                "type_total_co2e": round(et["type_total_co2e"], 4),
                "type_gas_totals": et["type_gas_totals"],
            })

        scope_total = round(scope_totals[scope_key], 4)
        if type_list or scope_total > 0:
            scope_data.append({
                "scope_key": scope_key,
                "scope_name": scope_display.get(scope_key, scope_key),
                "emission_types": type_list,
                "scope_total_co2e": scope_total,
            })

    total_co2e = round(sum(scope_totals.values()), 4)
    emission_type_labels = list(emission_type_totals_raw.keys())
    emission_type_values = [round(emission_type_totals_raw[k], 4) for k in emission_type_labels]
    record_count = len(records)
    device_count = len(devices)

    scope_labels = [scope_display[k] for k in scope_order if scope_totals[k] > 0]
    scope_values = [round(scope_totals[k], 4) for k in scope_order if scope_totals[k] > 0]

    # 各設備排放量（供圖表用）
    device_totals_raw: dict[str, float] = {}
    for scope_entry in scope_data:
        for et in scope_entry.get("emission_types", []):
            for dev in et.get("devices", []):
                name = dev.get("device_name", "未知")
                device_totals_raw[name] = dev.get("total_co2e", 0)
    sorted_devices = sorted(device_totals_raw.items(), key=lambda x: x[1], reverse=True)
    device_labels = [d[0] for d in sorted_devices[:10]]
    device_values = [round(d[1], 4) for d in sorted_devices[:10]]

    return templates.TemplateResponse(
        "result.html",
        {
            "request": request,
            "inventory_year": _get_target_year(session),
            "current_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "total_co2e": total_co2e,
            "record_count": record_count,
            "device_count": device_count,
            "scope_data": scope_data,
            "emission_type_labels": emission_type_labels,
            "emission_type_values": emission_type_values,
            "device_labels": device_labels,
            "device_values": device_values,
            "scope_labels": scope_labels,
            "scope_values": scope_values,
        },
    )


@router.get("/download-report-pdf")
async def download_report_pdf(session: Session = Depends(get_session)):
    from services.report_generator import create_report_draft
    year = _get_target_year(session)
    report = await create_report_draft(inventory_year=year)
    return RedirectResponse(url=f"/reports/{report.id}/pdf")


@router.post("/generate-ai-report")
async def generate_ai_report(
    request: Request,
    background_tasks: BackgroundTasks,
    session: Session = Depends(get_session),
):
    try:
        _resolve_llm_runtime_config()
    except RuntimeError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    account_id = request.session.get("user")
    snapshot = _build_report_snapshot(session, account_id=account_id)
    if snapshot["summary"]["record_count"] == 0:
        raise HTTPException(status_code=400, detail="目前沒有排放紀錄，無法生成報告。")

    task_id = str(uuid4())
    _set_task(
        task_id,
        status="pending",
        message="任務已建立，準備開始。",
        created_at=datetime.now().isoformat(),
        completed_at=None,
        file_path=None,
        error=None,
    )
    background_tasks.add_task(_run_ai_report_task, task_id, snapshot)

    return {
        "task_id": task_id,
        "status": "pending",
        "message": "任務已建立，請輪詢狀態。",
    }


@router.get("/ai-report-status/{task_id}")
async def get_ai_report_status(task_id: str):
    task = _get_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="找不到對應的報告任務。")

    return {
        "task_id": task_id,
        "status": task.get("status", "unknown"),
        "message": task.get("message", ""),
        "error": task.get("error"),
        "download_url": f"/result/download-ai-report/{task_id}"
        if task.get("status") == "completed"
        else None,
    }


@router.get("/download-ai-report/{task_id}")
async def download_ai_report(task_id: str):
    task = _get_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="找不到對應的報告任務。")
    if task.get("status") != "completed":
        raise HTTPException(status_code=409, detail="報告尚未完成，請稍後再試。")

    file_path_raw = task.get("file_path")
    if not file_path_raw:
        raise HTTPException(status_code=500, detail="報告檔案路徑不存在。")

    file_path = Path(file_path_raw)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="報告檔案不存在。")

    filename = f"ai_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return FileResponse(
        str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )
